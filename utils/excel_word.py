import io
import openpyxl
from docx import Document
from utils.data_helpers import tentar_converter_numero, limpar_texto_xml

def gerar_excel_bytes(dados_exportacao):
    output = io.BytesIO()
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Salvamento Remição"

    for item in dados_exportacao:
        ws.append([f"NOME: {item['nome']}"])
        ws.append([f"ORGANIZAÇÃO: {item['organiz']} | FUNÇÃO: {item['funcao']} | REMUNERAÇÃO: {item['remuneracao']} | SAÍDA: {item['saida']}"])
        ws.append([])

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            ws.append(headers)
            for idx_row, row_data in pivot_df.iterrows():
                row_vals = [tentar_converter_numero(idx_row)] + [tentar_converter_numero(v) for v in row_data.values]
                ws.append(row_vals)

        ws.append(["Total de Dias:", tentar_converter_numero(item['total_dias'])])
        ws.append([])
        ws.append([])

    wb.save(output)
    return output.getvalue()

def gerar_docx_bytes(dados_exportacao):
    output = io.BytesIO()
    doc = Document()
    doc.add_heading("Espaço de Dados para Salvamento", level=1)

    for item in dados_exportacao:
        doc.add_heading(limpar_texto_xml(f"NOME: {item['nome']}"), level=2)
        p_meta = doc.add_paragraph()
        p_meta.add_run(
            limpar_texto_xml(
                f"ORGANIZAÇÃO: {item['organiz']} | "
                f"FUNÇÃO: {item['funcao']} | "
                f"REMUNERAÇÃO: {item['remuneracao']} | "
                f"SAÍDA: {item['saida']}"
            )
        )

        pivot_df = item['pivot_df']
        if not pivot_df.empty:
            headers = ["ANO"] + list(pivot_df.columns)
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = limpar_texto_xml(h)

            for idx_row, row_data in pivot_df.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = limpar_texto_xml(idx_row)
                for i, val in enumerate(row_data.values):
                    row_cells[i+1].text = limpar_texto_xml(val)

        p_tot = doc.add_paragraph()
        p_tot.add_run(limpar_texto_xml(f"Total de Dias: {item['total_dias']}")).bold = True
        doc.add_paragraph()

    doc.save(output)
    return output.getvalue()
